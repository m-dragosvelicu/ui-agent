#!/usr/bin/env python3
"""Generate the blog post as a DOCX file."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Configure styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Building a UI/UX Research Agent for Web Development', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run('A practical guide to building an AI agent that automates design research')
run.italic = True

doc.add_paragraph()

# Introduction
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "If you've been doing web development for any amount of time, you know the cycle: you need a landing page, "
    "you reach for shadcn, you get something that looks like every other SaaS landing page from 2023. "
    "Or worse—you're stuck on a bug, you've been copy-pasting between ChatGPT and your codebase for an hour, "
    "and you're one \"I don't have access to your codebase\" away from throwing your laptop."
)
doc.add_paragraph(
    "The naive workflow looks like this: complain to ChatGPT about what's broken, ask it what info it needs, "
    "switch to another tool to extract the relevant code, paste everything back, and repeat until sanity depletes. "
    "Too much context-switching. Too many iterations."
)
doc.add_paragraph(
    "What if we could externalize this entire process? An agent that takes your complaint, digs through your codebase, "
    "researches solutions, and comes back with actual fresh ideas—not the same recycled Tailwind templates."
)
doc.add_paragraph("That's what we're building today: fixitmany.")

# Use Cases
doc.add_heading('Use Cases', level=1)

doc.add_heading('Use Case 1: Design Inspiration', level=2)
doc.add_paragraph(
    "You want to improve a landing page but you're out of ideas. Or you have a vision, but every AI suggestion "
    "is the same hero-section-with-gradient nonsense. The agent should:"
)
bullets = doc.add_paragraph()
bullets.add_run("• Understand your current design\n")
bullets.add_run("• Research fresh, unconventional approaches\n")
bullets.add_run("• Suggest specific implementations that aren't just \"add a CTA button\"")

doc.add_heading('Use Case 2: Bug Research', level=2)
doc.add_paragraph("You've got a bug you can't crack. The agent should:")
bullets2 = doc.add_paragraph()
bullets2.add_run("• Extract relevant code segments automatically\n")
bullets2.add_run("• Research similar issues and solutions\n")
bullets2.add_run("• Come back with targeted fixes, not generic advice")

doc.add_heading('Use Case 3: Component Modernization', level=2)
doc.add_paragraph(
    "You have an existing component that works but feels dated or clunky. Maybe it's a card, a modal, or a data table. The agent should:"
)
bullets3 = doc.add_paragraph()
bullets3.add_run("• Analyze your current implementation\n")
bullets3.add_run("• Research modern patterns for that component type\n")
bullets3.add_run("• Suggest specific improvements (accessibility, animations, UX patterns)\n")
bullets3.add_run("• Provide updated code that you can drop in")

# Architecture
doc.add_heading('Architecture', level=1)
doc.add_paragraph(
    "We're building a tool-using agent with five core capabilities:"
)
arch = doc.add_paragraph()
arch.add_run("1. read_file").bold = True
arch.add_run(" - Extract code from your project\n")
arch.add_run("2. list_files").bold = True
arch.add_run(" - Understand project structure\n")
arch.add_run("3. write_file").bold = True
arch.add_run(" - Output improved components\n")
arch.add_run("4. search_web").bold = True
arch.add_run(" - Research design trends and solutions\n")
arch.add_run("5. fetch_url").bold = True
arch.add_run(" - Read documentation and articles")

doc.add_paragraph(
    "The key insight: the agent decides which tools to use and when. You give it a problem, it figures out the research path. "
    "The agentic loop follows a simple pattern: Think → Act → Observe → Repeat."
)

doc.add_paragraph(
    "The agent supports multiple LLM providers: Anthropic (Claude), OpenAI (GPT), and Google (Gemini). "
    "This flexibility allows you to choose the model that best fits your needs or switch between providers as needed."
)

# Step-by-Step Tutorial
doc.add_heading('Step-by-Step Tutorial', level=1)

doc.add_heading('Step 1: Project Setup', level=2)
doc.add_paragraph("Clone the repository and install as a global CLI command:")
code1 = doc.add_paragraph()
code1.style = 'No Spacing'
run = code1.add_run(
    "git clone https://github.com/m-dragosvelicu/ui-agent.git\n"
    "cd ui-agent/uiux-agent\n"
    "pip install -e ."
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph("This installs fixitmany as a command you can run from anywhere.")

doc.add_heading('Step 2: Configure API Keys', level=2)
doc.add_paragraph(
    "Set your API key as an environment variable. You only need the key for the provider you want to use. "
    "Gemini is the default provider:"
)
code_env = doc.add_paragraph()
code_env.style = 'No Spacing'
run = code_env.add_run(
    "# For Gemini (default)\n"
    "export GOOGLE_API_KEY=\"your-key\"\n\n"
    "# For Claude\n"
    "export ANTHROPIC_API_KEY=\"your-key\"\n\n"
    "# For OpenAI\n"
    "export OPENAI_API_KEY=\"your-key\""
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph("Or create a .env file in the uiux-agent directory:")
code_env2 = doc.add_paragraph()
code_env2.style = 'No Spacing'
run = code_env2.add_run("cp .env.example .env\n# Edit .env and add your key(s)")
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_heading('Step 3: Run the Agent', level=2)
doc.add_paragraph("Point it at your project and describe what you need:")
code2 = doc.add_paragraph()
code2.style = 'No Spacing'
run = code2.add_run(
    "# Basic usage (Gemini is default)\n"
    "fixitmany \"Improve my landing page\" -p ./my-project\n\n"
    "# Use Claude instead\n"
    "fixitmany --provider anthropic \"Fix this bug\" -p ./my-project\n\n"
    "# Use OpenAI instead\n"
    "fixitmany --provider openai \"Modernize this component\" -p ./my-project\n\n"
    "# Interactive chat mode\n"
    "fixitmany --interactive -p ./my-project"
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

# How It Works
doc.add_heading('How It Works Under the Hood', level=1)

doc.add_heading('The Tools', level=2)
doc.add_paragraph(
    "The agent has five tools defined in tools.py. Each tool is a Python function that performs a specific action:"
)
bullets4 = doc.add_paragraph()
bullets4.add_run("• read_file: Reads file contents with truncation for large files\n")
bullets4.add_run("• list_files: Recursively lists files, skipping node_modules and .git\n")
bullets4.add_run("• write_file: Creates directories as needed and writes content\n")
bullets4.add_run("• search_web: Uses DuckDuckGo HTML search for web results\n")
bullets4.add_run("• fetch_url: Extracts clean text content from web pages")

doc.add_heading('The Provider Abstraction', level=2)
doc.add_paragraph(
    "The providers.py file defines a common interface that works with Anthropic, OpenAI, and Gemini. "
    "Each provider converts messages and tool definitions to its native format, allowing you to switch "
    "between models with a simple --provider flag."
)

doc.add_heading('The Agent Loop', level=2)
doc.add_paragraph(
    "The agent.py file contains the main loop. It continuously calls the LLM until it gets a final response "
    "or reaches the maximum iteration limit (15 by default)."
)
doc.add_paragraph("The loop works as follows:")
bullets5 = doc.add_paragraph()
bullets5.add_run("1. Send the user's task to the LLM with available tools\n")
bullets5.add_run("2. If the LLM returns tool calls, execute them\n")
bullets5.add_run("3. Feed the tool results back to the LLM\n")
bullets5.add_run("4. Repeat until the LLM provides a final answer\n")
bullets5.add_run("5. Return the final text response to the user")

# Real Example
doc.add_heading('Real Example: Improving a Product Card', level=1)
doc.add_paragraph(
    "Let's see the agent in action. We have a basic product card component—functional but forgettable. "
    "It's the kind of card you've seen on every e-commerce tutorial since 2019: border, rounded corners, shadow, done."
)

doc.add_paragraph("The original component:")
code3 = doc.add_paragraph()
code3.style = 'No Spacing'
run = code3.add_run(
    "export function ProductCard({ title, price, image, description, onAddToCart }) {\n"
    "  return (\n"
    "    <div className=\"border rounded-lg p-4 shadow-sm\">\n"
    "      <img src={image} alt={title} className=\"w-full h-48 object-cover rounded\" />\n"
    "      <h3 className=\"text-lg font-semibold mt-2\">{title}</h3>\n"
    "      <p className=\"text-gray-600 text-sm mt-1\">{description}</p>\n"
    "      <div className=\"flex justify-between items-center mt-4\">\n"
    "        <span className=\"text-xl font-bold\">${price}</span>\n"
    "        <button onClick={onAddToCart} className=\"bg-blue-500 text-white px-4 py-2 rounded\">\n"
    "          Add to Cart\n"
    "        </button>\n"
    "      </div>\n"
    "    </div>\n"
    "  );\n"
    "}"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph("Running the agent:")
task_p = doc.add_paragraph()
task_p.style = 'No Spacing'
run = task_p.add_run(
    "fixitmany \"This product card is boring. Every e-commerce site looks like this. "
    "I want something that feels premium and modern - maybe some micro-interactions, "
    "better visual hierarchy.\" -p ./example-project"
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_heading('What the Agent Does', level=2)
doc.add_paragraph("The agent autonomously executes the following steps:")
steps = doc.add_paragraph()
steps.add_run("1. ").bold = True
steps.add_run("list_files(\"./example-project\") → maps the project structure\n")
steps.add_run("2. ").bold = True
steps.add_run("read_file(\"./example-project/components/ProductCard.tsx\") → examines current code\n")
steps.add_run("3. ").bold = True
steps.add_run("search_web(\"modern product card design trends 2024 e-commerce\") → finds current patterns\n")
steps.add_run("4. ").bold = True
steps.add_run("search_web(\"framer motion product card hover animations\") → digs into specific techniques\n")
steps.add_run("5. ").bold = True
steps.add_run("write_file(\"./example-project/components/ProductCard.improved.tsx\") → outputs the improved version")

doc.add_heading('Agent Output', level=2)
doc.add_paragraph("The agent's analysis:")
quote = doc.add_paragraph()
quote.style = 'Quote'
quote.add_run(
    "\"This card screams 'Bootstrap tutorial from 2019'. Here's what's wrong: Zero hierarchy—the price and button "
    "compete for attention. No interactivity—static as a newspaper ad. Generic shadow—the 'I learned CSS yesterday' special.\""
)

doc.add_paragraph("The agent recommends:")
rec = doc.add_paragraph()
rec.add_run("• Stacked visual layers").bold = True
rec.add_run(" using subtle transforms on hover\n")
rec.add_run("• Price badge").bold = True
rec.add_run(" that floats over the image corner\n")
rec.add_run("• Animated cart button").bold = True
rec.add_run(" that expands with a satisfying spring\n")
rec.add_run("• Quick-view overlay").bold = True
rec.add_run(" on image hover")

doc.add_paragraph(
    "The agent then writes a complete improved component using Framer Motion for animations, "
    "with proper hover states, visual hierarchy, and micro-interactions."
)

# What Makes This Different
doc.add_heading('What Makes This Different', level=1)
diff = doc.add_paragraph()
diff.add_run("1. Externalized thinking").bold = True
diff.add_run(" – You fire off the task and come back to results\n")
diff.add_run("2. Tool-use autonomy").bold = True
diff.add_run(" – The agent decides what to research, not you\n")
diff.add_run("3. Fresh perspectives").bold = True
diff.add_run(" – By searching beyond its training data, it finds current trends\n")
diff.add_run("4. Codebase awareness").bold = True
diff.add_run(" – It reads your actual code, not generic examples\n")
diff.add_run("5. Multi-provider support").bold = True
diff.add_run(" – Switch between Claude, GPT, or Gemini based on your needs")

# Quick Reference
doc.add_heading('Quick Reference', level=1)
doc.add_paragraph("Once installed, use fixitmany from anywhere:")
ref = doc.add_paragraph()
ref.style = 'No Spacing'
run = ref.add_run(
    "fixitmany \"task\" -p <path>              # Basic usage (Gemini default)\n"
    "fixitmany --provider anthropic \"task\"   # Use Claude instead\n"
    "fixitmany --provider openai \"task\"      # Use GPT instead\n"
    "fixitmany -q \"task\" -p <path>           # Quiet mode (less output)\n"
    "fixitmany --interactive -p <path>       # Chat mode\n"
    "fixitmany --model gemini-2.5-pro \"task\" # Custom model\n"
    "fixitmany --help                        # Show all options"
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

# Safety
doc.add_heading('Safety', level=1)
doc.add_paragraph(
    "The agent NEVER overwrites existing files. If you ask it to write to a file that already exists, "
    "it creates a new file with a .new suffix instead. For example, if ProductCard.tsx exists, "
    "the agent will create ProductCard.new.tsx instead of overwriting the original."
)

# Next Steps
doc.add_heading('Next Steps', level=1)
doc.add_paragraph("This is a starting point. To make it production-ready:")
next_p = doc.add_paragraph()
next_p.add_run("• Add more tools (screenshot analysis, Figma integration, performance checks)\n")
next_p.add_run("• Implement caching for repeated searches\n")
next_p.add_run("• Add a simple UI (Streamlit works great)\n")
next_p.add_run("• Fine-tune the system prompt for your specific stack")

doc.add_paragraph(
    "The goal isn't to replace your judgment—it's to do the research grunt work so you can focus on "
    "the creative decisions that actually matter."
)

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "Stop settling for stale components. Build agents that find the fresh stuff for you. "
    "The key difference from manual copy-paste workflows: the agent makes the research decisions itself. "
    "It doesn't just search once—it iterates based on what it finds, then writes the improved component directly to your project."
)

doc.add_paragraph()
final = doc.add_paragraph()
final.add_run("Repository: ").bold = True
final.add_run("https://github.com/m-dragosvelicu/ui-agent")

# Save
doc.save('/Users/dragosvelicu/Documents/University RAU/ai-agent-article/Building_a_UI_UX_Research_Agent.docx')
print("Created: Building_a_UI_UX_Research_Agent.docx")
